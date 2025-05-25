import os
import tempfile
from io import BytesIO
import streamlit as st
from dotenv import load_dotenv
from langchain.document_loaders import PyPDFLoader
from langchain.docstore.document import Document
from langchain.text_splitter import TokenTextSplitter
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains.summarize import load_summarize_chain
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from docx import Document as WordDocument  # pip install python-docx

# Load API key
load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

# Prompt templates
prompt_template = "Write possible interview questions from the following text:\n\n{text}\n\nQuestions:"
refine_template = ("We have some existing interview questions: {existing_answer}\n\n"
                   "Based on the additional context below, refine and add more questions:\n\n{text}\n\nRefined Questions:")

# PDF processing function
def file_processing(file_path):
    loader = PyPDFLoader(file_path)
    data = loader.load()

    content = ''.join([page.page_content for page in data])

    splitter_ques_gen = TokenTextSplitter(
        model_name='gpt-3.5-turbo',
        chunk_size=10000,
        chunk_overlap=200
    )
    document_ques_gen = [Document(page_content=t) for t in splitter_ques_gen.split_text(content)]

    splitter_ans_gen = TokenTextSplitter(
        model_name='gpt-3.5-turbo',
        chunk_size=1000,
        chunk_overlap=100
    )
    document_answer_gen = splitter_ans_gen.split_documents(document_ques_gen)

    return document_ques_gen, document_answer_gen

# LLM pipeline
def llm_pipeline(file_path):
    document_ques_gen, document_answer_gen = file_processing(file_path)

    llm_ques_gen_pipeline = ChatOpenAI(temperature=0.3, model="gpt-3.5-turbo")
    PROMPT_QUESTIONS = PromptTemplate(template=prompt_template, input_variables=["text"])
    REFINE_PROMPT_QUESTIONS = PromptTemplate(template=refine_template, input_variables=["existing_answer", "text"])

    ques_gen_chain = load_summarize_chain(
        llm=llm_ques_gen_pipeline,
        chain_type="refine",
        verbose=True,
        question_prompt=PROMPT_QUESTIONS,
        refine_prompt=REFINE_PROMPT_QUESTIONS
    )

    ques = ques_gen_chain.run(document_ques_gen)
    ques_list = [q.strip() for q in ques.split("\n") if q.strip().endswith(("?", "."))]

    embeddings = OpenAIEmbeddings()
    vector_store = FAISS.from_documents(document_answer_gen, embeddings)

    answer_generation_chain = RetrievalQA.from_chain_type(
        llm=ChatOpenAI(temperature=0.1, model="gpt-3.5-turbo"),
        chain_type="stuff",
        retriever=vector_store.as_retriever()
    )

    return answer_generation_chain, ques_list

# Function to generate Word document from Q&A pairs
def generate_word_doc(questions, answers):
    doc = WordDocument()
    doc.add_heading("Interview Questions & Answers", level=1)

    for idx, (q, a) in enumerate(zip(questions, answers), 1):
        doc.add_paragraph(f"Q{idx}: {q}", style='List Number')
        doc.add_paragraph(f"A{idx}: {a}")
        doc.add_paragraph()  # blank line

    # Save to in-memory bytes buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.set_page_config(page_title="PDF Q&A Generator", page_icon="📄")
st.title("📄 Interview Question & Answer Generator from PDF")

uploaded_file = st.file_uploader("Upload a PDF file (Max 5 pages)", type=["pdf"])

if uploaded_file is not None:
    st.success(f"✅ Uploaded file: {uploaded_file.name}")

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name

        st.info("⏳ Processing your file...")
        qa_chain, questions = llm_pipeline(tmp_path)

        answers = []
        for q in questions:
            answer = qa_chain.run(q)
            answers.append(answer)

        st.success("✅ Questions and answers generated!")

        # Display Q&A interleaved
        for idx, (q, a) in enumerate(zip(questions, answers), 1):
            st.markdown(f"**Q{idx}: {q}**")
            st.markdown(f"🟢 **A:** {a}")
            st.markdown("---")

        # Generate and offer download of Word doc
        word_buffer = generate_word_doc(questions, answers)
        st.download_button(
            label="📥 Download Q&A as Word Document",
            data=word_buffer,
            file_name="interview_qa.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        os.remove(tmp_path)

    except Exception as e:
        st.error(f"❌ Error while processing PDF: {e}")

else:
    st.info("📥 Upload a PDF to begin.")
